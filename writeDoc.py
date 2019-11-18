#!/usr/bin/env python
#-*- coding:utf-8 -*-

from docxtpl import DocxTemplate
import os
from docx.shared import Pt, Inches, Cm, Length
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import docx
import datetime

def getDir():
    return os.getcwd()


class doc(object):
    def __init__(self):
        base_url = getDir() + '/user/user_template/'            #初始化生成一个doc对象
        asset_url = base_url + 'demo.docx'
        self.tpl = DocxTemplate(asset_url)
        self.webvul = self.tpl.new_subdoc()                #web应用漏洞
        self.appvul = self.tpl.new_subdoc()                #应用程序漏洞
        self.devicevul = self.tpl.new_subdoc()             #网络设备漏洞
        self.sysvul = self.tpl.new_subdoc()                #操作系统应用漏洞
        self.time = self.tpl.new_subdoc()
        self.number = 0

    def add_title(self, title, type , level=3):
        doc = self.check_type(type)
        run = doc.add_heading('', level=level).add_run(title)
        run.font.name = "宋体"
        run.font.size = Pt(14)
        r = run.element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    def add_table(self,jsons, type=1):
        '''
        在doc中添加表格
        :param json:
                type:1为web漏洞2为应用程序漏洞,3为网络设备漏洞4为系统漏洞
        :return:
        '''

        #获取要写入的subdoc节点
        doc = self.check_type(type)

        for json in jsons:
            # j控制json字典下标
            #print(json)
            j = 0
            title = json['title']                                 #生成标题
            self.number += 1
            self.add_title(title,type=type)
            table = doc.add_table(rows=len(json)-1, cols=2, style="Style1")#新建一个表格
            table.autofit = False
            for key, value in json.items():
                if j == 0:
                    j +=1
                else:
                    table.columns[0].width = Cm(3)
                    table.columns[1].width = Cm(12)
                    table.cell(j - 1, 0).width = Cm(3)   #设置单元格宽度
                    table.cell(j - 1, 1).width = Cm(12)
                    #table.alignment=WD_TABLE_ALIGNMENT.RIGHT 设置对齐方式

                    keyCell = table.cell(j-1, 0)            #表格赋值
                    valueCell = table.cell(j-1, 1)

                    #设置key单元格字体与字体大小
                    key_paragraph = keyCell.paragraphs[0]
                    #keyRun = keyCell.paragraphs[0].add_run(key)
                    keyRun = key_paragraph.add_run(key)
                    keyRun.font.name = u'微软雅黑'  # 设置字体
                    keyRun._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
                    keyRun.font.size = Pt(10.5)  # 设置字号为五号
                    key_paragraph.paragraph_format.line_spacing = WD_LINE_SPACING.ONE_POINT_FIVE#设置1.5倍行间距
                    key_paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT #设置水平对齐方式右对齐
                    keyCell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                    #设置value单元格字体与字体大小
                    val_paragraph = valueCell.paragraphs[0]

                    #valueRun = valueCell.paragraphs[0].add_run(value)  # 填入的内容

                    valueRun = val_paragraph.add_run(value)
                    valueRun.font.name = u'微软雅黑'  # 设置字体
                    valueRun.font.size = Pt(10.5)  # 设置字号为五号
                    valueRun._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
                    val_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE#设置1.5倍行间距
                    val_paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT#设置水平对齐方式左对齐
                    valueCell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    #keyCell.text = key
                    #valueCell.text = value
                    #print(key,value)
                    j = j+1
        #赋值给初始话的subdoc值
        if type==1:
            self.webvul=doc
        elif type==2:
            self.appvul= doc
        elif type==3:
            self.devicevul = doc
        else:
            self.sysvul = doc

    #保存docx文档
    def save_doc(self,current_time, start_time):
        #添加doc编辑时间
        filename = "上海驭胜信息安全通告（%s至%s).docx" % (str(start_time),str(current_time))
        content = {'subdoc': self.webvul,
                        'appdoc':self.appvul,
                        'devicedoc': self.devicevul,
                        'sysdoc': self.sysvul,
                        'time': current_time ,
                        'start_time': start_time,
                        'end_time': current_time}
        self.tpl.render(content)
        self.tpl.save(filename)
        print("file %s success to save!" % filename)

    def check_type(self,num):
        if num==1:
            return self.webvul
        elif num ==2:
            return self.appvul
        elif num ==3:
            return self.devicevul
        else:
            return self.sysvul

    def add_hyperlink(self,paragraph, text, url):
        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        # Create a w:r element and a new w:rPr element
        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        # Create a new Run object and add the hyperlink into it
        r = paragraph.add_run ()
        r.font.name = u'微软雅黑'  # 设置字体
        r.font.size = Pt(10.5)  # 设置字号为五号
        r._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        r._r.append (hyperlink)

        # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
        # Delete this if using a template that has the hyperlink style in it
        r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        r.font.underline = True

        return hyperlink